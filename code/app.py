import streamlit as st
import os
import io
import re
import unicodedata
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 0. XỬ LÝ ĐƯỜNG DẪN (CHỐNG LỖI FILE NOT FOUND) =====
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_BIA_PATH = os.path.join(CURRENT_DIR, "Bia.docx")

# ===== 1. DATABASE 9 HỌC PHẦN (CHUẨN EXCEL) =====
DATA_BTL = {
    "HP_01: Giáo dục học đại cương": {
        "de_tai": "PHÂN TÍCH CÁC CHỨC NĂNG XÃ HỘI CỦA GIÁO DỤC VÀ LIÊN HỆ THỰC TIỄN",
        "cau_hoi": "Bằng lý luận và thực tiễn, anh (chị) hãy phân tích các chức năng xã hội của giáo dục. Từ đó, anh (chị) hãy liên hệ với việc thực hiện các chức năng này ở Việt Nam hiện nay.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_02: Sử dụng phương tiện kỹ thuật KH trong dạy học": {
        "de_tai": "PHÂN TÍCH CẤU TẠO, NGUYÊN LÝ HOẠT ĐỘNG CỦA CÁC PHƯƠNG TIỆN KỸ THUẬT HIỆN ĐẠI",
        "cau_hoi": "Anh (Chị) hãy trình bày cấu tạo, nguyên lý hoạt động và quy trình sử dụng của một số phương tiện kỹ thuật dạy học hiện đại (máy chiếu, bảng tương tác...). Phân tích ưu nhược điểm khi vận dụng vào môn học cụ thể.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_03: Lý luận dạy học đại học": {
        "de_tai": "VẬN DỤNG CÁC NGUYÊN TẮC DẠY HỌC HIỆN ĐẠI TRONG ĐÀO TẠO ĐẠI HỌC",
        "cau_hoi": "Phân tích hệ thống các nguyên tắc dạy học đại học. Đề xuất phương án vận dụng các nguyên tắc này vào giảng dạy một học phần cụ thể thuộc chuyên ngành anh (chị) đang theo đuổi.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.3, "font_sz": 14
    },
    "HP_04: Tâm lý học đại học": {
        "de_tai": "PHÂN TÍCH CÁC GIAI ĐOẠN CỦA QUÁ TRÌNH TRI GIÁC VÀ ỨNG DỤNG TRONG DẠY HỌC",
        "cau_hoi": "Trình bày các giai đoạn của quá trình tri giác. Phân tích các đặc điểm tâm lý của sinh viên đại học và đề xuất các biện pháp sư phạm nhằm phát huy tính tích cực nhận thức của người học.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_05: Đánh giá trong giáo dục đại học": {
        "de_tai": "PHÂN TÍCH VAI TRÒ CỦA ĐÁNH GIÁ TRONG ĐỔI MỚI GIÁO DỤC ĐẠI HỌC",
        "cau_hoi": "Trình bày vai trò, chức năng của kiểm tra đánh giá trong giáo dục đại học. Phân tích thực trạng và đề xuất giải pháp đổi mới hình thức đánh giá theo định hướng phát triển năng lực người học.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_06: Quản lý nhà nước về giáo dục và đào tạo": {
        "de_tai": "PHÂN TÍCH CÁC NGUYÊN TẮC QUẢN LÝ NHÀ NƯỚC VỀ GIÁO DỤC",
        "cau_hoi": "Trình bày các nguyên tắc và nội dung quản lý nhà nước về giáo dục và đào tạo. Liên hệ thực hiện tại một cơ sở giáo dục cụ thể và đề xuất giải pháp hoàn thiện.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_07: Phát triển chương trình đào tạo đại học": {
        "de_tai": "PHÂN TÍCH CÁC THÀNH PHẦN CỦA CHƯƠNG TRÌNH ĐÀO TẠO THEO TIẾP CẬN CDIO",
        "cau_hoi": "Trình bày các thành phần cơ bản của một chương trình đào tạo. Phân tích quy trình thiết kế chương trình đào tạo theo tiếp cận năng lực và xây dựng chuẩn đầu ra cho một học phần cụ thể.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_08: Phương pháp nghiên cứu khoa học giáo dục": {
        "de_tai": "PHÂN TÍCH CÁC PHƯƠNG PHÁP NGHIÊN CỨU KHOA HỌC TRONG GIÁO DỤC",
        "cau_hoi": "Trình bày quy trình thực hiện một đề tài nghiên cứu khoa học giáo dục. Phân tích các phương pháp thu thập thông tin (phỏng vấn, bảng hỏi, quan sát) và lập kế hoạch nghiên cứu cho một đề tài tự chọn.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    },
    "HP_09: Kỹ năng mềm": {
        "de_tai": "PHÂN TÍCH VAI TRÒ CỦA KỸ NĂNG GIAO TIẾP TRONG HOẠT ĐỘNG SƯ PHẠM",
        "cau_hoi": "Trình bày tầm quan trọng của kỹ năng mềm đối với giảng viên đại học. Phân tích các kỹ năng giao tiếp, kỹ năng làm việc nhóm và kỹ năng giải quyết vấn đề trong môi trường giáo dục hiện đại.",
        "le": (2.0, 2.0, 3.0, 2.0), "spacing": 1.5, "font_sz": 14
    }
}

# ===== 2. TIỆN ÍCH HỆ THỐNG =====
def remove_vietnamese_accent(s):
    s = s.replace("Đ", "D").replace("đ", "d")
    s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode('ascii')
    return re.sub(r'[^a-zA-Z0-9_]', '_', s)

def add_page_number(paragraph, alignment_code):
    if "LEFT" in alignment_code: paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif "RIGHT" in alignment_code: paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin'); run._r.append(fld1)
    instr = OxmlElement('w:instrText'); instr.text = "PAGE"; run._r.append(instr)
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end'); run._r.append(fld2)

def set_font_style(run, font_name="Times New Roman", size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name); rFonts.set(qn('w:hAnsi'), font_name); rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def strictly_clean_content(text):
    ai_patterns = [r"Dưới đây là.*:", r"Đoạn văn đã được chỉnh sửa.*:", r"Chắc chắn rồi.*:"]
    for pattern in ai_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.DOTALL)
    return re.sub(r'[*#_~-]', '', text).strip()

# ===== 3. GIAO DIỆN STREAMLIT =====
st.set_page_config(page_title="AI Report Generator Pro", layout="wide")

with st.sidebar:
    st.header("⚙️ Cấu hình AI")
    provider = st.selectbox("Provider", ["Groq", "Gemini"])
    model_choice = st.selectbox("Model", ["llama-3.3-70b-versatile", "qwen-2.5-32b"] if provider == "Groq" else ["gemini-1.5-pro", "gemini-2.0-flash"])
    api_key_input = st.text_input("API Key", type="password")

    st.divider()
    st.header("📏 Thông số Trang (Có thể chỉnh)")
    selected_hp = st.selectbox("Chọn Mã Học Phần", list(DATA_BTL.keys()))
    preset = DATA_BTL[selected_hp]
    
    c1, c2 = st.columns(2)
    with c1:
        m_top = st.number_input("Trên (cm)", 0.0, 5.0, preset["le"][0])
        m_left = st.number_input("Trái (cm)", 0.0, 5.0, preset["le"][2])
    with c2:
        m_bottom = st.number_input("Dưới (cm)", 0.0, 5.0, preset["le"][1])
        m_right = st.number_input("Phải (cm)", 0.0, 5.0, preset["le"][3])

    f_sz = st.number_input("Cỡ chữ", 10, 16, preset["font_sz"])
    f_sp = st.number_input("Dãn dòng", 1.0, 2.5, preset["spacing"])
    
    st.header("🔢 Đánh số trang")
    pg_align = st.selectbox("Vị trí ngang", ["RIGHT", "LEFT", "CENTER"], index=0)

    # --- KHU VỰC FIX LỖI FILE BÌA ---
    st.divider()
    st.header("📁 Quản lý File Bìa")
    if not os.path.exists(DEFAULT_BIA_PATH):
        st.error("⚠️ Không tìm thấy file Bia.docx tự động.")
        uploaded_bia = st.file_uploader("Vui lòng tải file Bia.docx lên đây để dùng:", type=["docx"])
    else:
        st.success("✅ Đã tìm thấy Bia.docx trong thư mục.")
        uploaded_bia = None

st.title("🎓 Hệ thống Tạo Tiểu luận Ultimate (VNU-HPU2 Edition)")

col_in1, col_in2 = st.columns([1, 2])
with col_in1:
    st.subheader("👤 Thông tin Bìa")
    ten_hv = st.text_input("Học viên", "Đặng Nhật Minh")
    sbd = st.text_input("SBD", "39")
    mon_hoc = st.text_input("Môn học", selected_hp.split(": ")[1])
    de_tai = st.text_area("Tên đề tài (In bìa)", preset["de_tai"], height=100)

with col_in2:
    st.subheader("🤖 Yêu cầu nội dung")
    yeu_cau = st.text_area("Câu hỏi chi tiết", preset["cau_hoi"], height=220)

# ===== 4. LOGIC AI & XUẤT FILE =====
def call_ai(key, provider, prompt, model):
    sys = "Bạn là Tiến sĩ Giáo dục. CHỈ xuất nội dung tiểu luận. KHÔNG dẫn dắt, KHÔNG Markdown."
    try:
        if provider == "Gemini":
            import google.generativeai as genai
            genai.configure(api_key=key)
            return genai.GenerativeModel(model).generate_content(f"{sys}\n\n{prompt}").text
        else:
            from groq import Groq
            client = Groq(api_key=key)
            res = client.chat.completions.create(model=model, messages=[{"role": "system", "content": sys}, {"role": "user", "content": prompt}])
            return res.choices[0].message.content
    except Exception as e: return f"Lỗi: {e}"

if st.button("🚀 XUẤT TIỂU LUẬN HOÀN THIỆN"):
    key = api_key_input or st.secrets.get(f"{provider.upper()}_API_KEY")
    if not key: st.error("Thiếu Key!"); st.stop()
    
    # Xác định nguồn file bìa (Ưu tiên file upload, sau đó tới file cục bộ)
    if uploaded_bia is not None:
        doc = Document(uploaded_bia)
    elif os.path.exists(DEFAULT_BIA_PATH):
        doc = Document(DEFAULT_BIA_PATH)
    else:
        st.error("❌ Không có file Bia.docx để thực hiện! Hãy kiểm tra folder hoặc tải lên ở Sidebar.")
        st.stop()

    status = st.empty(); prog = st.progress(0)
    
    # 1. AI viết bài
    outline = call_ai(key, provider, f"Lập dàn ý 6 mục lớn cho tiểu luận: {de_tai}", model_choice)
    sections = [s.strip() for s in outline.split('\n') if len(s.strip()) > 10][:6]

    final_content = []
    for i, sec in enumerate(sections):
        status.info(f"✍️ Đang soạn Chương {i+1}: {sec}")
        draft = call_ai(key, provider, f"Viết 1000 từ cho mục '{sec}' của đề tài '{de_tai}'. Bám sát câu hỏi: {yeu_cau}", model_choice)
        polished = call_ai(key, provider, f"Biên tập lại: Xóa dấu AI, sửa câu tự nhiên. CẤM dẫn dắt: {draft}", model_choice)
        final_content.append((sec, strictly_clean_content(polished)))
        prog.progress(int((i+1)/len(sections)*100))

    # 2. Xử lý Word
    for p in doc.paragraphs:
        maps = {"{{HO_TEN}}": ten_hv, "{{SBD}}": sbd, "{{MON_HOC}}": mon_hoc.upper(), "{{TEN_DE_TAI}}": de_tai.upper()}
        for k, v in maps.items():
            if k in p.text:
                for run in p.runs:
                    if k in run.text: run.text = run.text.replace(k, str(v)); set_font_style(run)

    # 3. Mục lục & Định dạng lề
    doc.add_page_break()
    p_toc = doc.add_paragraph(); p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font_style(p_toc.add_run("MỤC LỤC"), size=16, bold=True)
    run_toc = doc.add_paragraph().add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin'); run_toc._r.append(fld1)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_toc._r.append(instr); fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate'); run_toc._r.append(fld2); fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end'); run_toc._r.append(fld3)

    doc.add_page_break()
    sec_word = doc.sections[-1]
    sec_word.top_margin, sec_word.bottom_margin = Cm(m_top), Cm(m_bottom)
    sec_word.left_margin, sec_word.right_margin = Cm(m_left), Cm(m_right)
    
    doc.sections[0].different_first_page_header_footer = True
    add_page_number(sec_word.footer.paragraphs[0], pg_align)

    for idx, (title, text) in enumerate(final_content):
        h = doc.add_paragraph(style='Heading 1')
        run_h = h.add_run(f"{idx + 1}. {title.upper()}")
        set_font_style(run_h, size=f_sz, bold=True)
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = f_sp
                run_p = p.runs[0] if p.runs else p.add_run()
                set_font_style(run_p, size=f_sz)

    file_final = f"BTL_{selected_hp.split(':')[0]}_{remove_vietnamese_accent(ten_hv)}.docx"
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    status.success(f"✅ Đã tạo xong bài làm!")
    st.download_button(label=f"📥 Tải xuống bài làm", data=buf, file_name=file_final)